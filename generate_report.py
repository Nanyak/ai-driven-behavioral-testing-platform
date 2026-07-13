import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_FILE = "VDT26_Dinh_Quyet_Thang_Bao_cao.docx"
FIGURE_DIR = Path("assets/report")
REPORT_TITLE = "BÁO CÁO MINI-PROJECT"
TOPIC = "Hệ thống tự động hóa kiểm thử Backend dựa trên phân tích hành vi người dùng"
AUTHOR = "Đinh Quyết Thắng"
EMAIL = "dinhquyetthang1303@gmail.com"
PROGRAM = "Viettel Digital Talent 2026"
FIELD = "Software Engineer"
MENTOR = "Đoàn Chí Tùng"
UNIT = "Viettel Digital Service (VDS)"


PROJECT_SNAPSHOT = {
    "latest_report": {
        "run_id": "run-2026-07-07-104122",
        "generated_at": "2026-07-07T10:41:22.630Z",
        "status": "green",
        "executed": 9,
        "passed": 9,
        "failed": 0,
        "skipped": 0,
    },
    "candidate_run": {
        "run_id": "2026-07-07T01-30-53-137Z",
        "candidate_count": 18,
        "min_support": 3,
        "prefixspan_patterns": 127764,
        "ngram_patterns": 189,
        "per_persona": {
            "guest_shopper": 3,
            "registered_customer": 10,
            "admin_operator": 5,
        },
        "holdout_support": 6,
        "holdout_floor": 6,
        "negative_control_passes": True,
        "macro_f1_cart_read": 0.9646,
        "accuracy_cart_read": 0.9547,
    },
    "postgres": {
        "decisions": 9,
        "approved_decisions": 9,
        "run_index": 11,
        "manifest": 10,
        "invariants": 295,
        "verified_invariants": 113,
        "unverified_invariants": 182,
    },
    "minio": {
        "approved_specs": 9,
        "candidates": 1,
        "endpoint_behavior_docs": 6,
        "goldens": 56,
        "reports": 33,
        "sessions": 1,
        "specs": 29,
        "validation": 1,
    },
    "mutation_eval": {
        "generated_at": "2026-07-07T18:39:25.188Z",
        "target": "all",
        "total_mutants": 150,
        "killed": 40,
        "survived": 5,
        "inconclusive": 105,
        "mutation_score": 0.8889,
        "executability_rate": 1.0,
        "baseline_clean": True,
        "note": "Mutation score được tính trên 45 mutant đo được; 105 mutant inconclusive chủ yếu do không được áp dụng hoặc chưa được luồng baseline exercise.",
    },
}


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, align="justify", first_line=True, spacing_after=6):
    alignments = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0.35) if first_line and align == "justify" else None


def add_paragraph(doc, text, bold=False, italic=False, align="justify", size=13):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align=align)
    for index, line in enumerate(str(text).splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level, page_break_before=False):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align="left", first_line=False, spacing_after=6)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.keep_with_next = True
    # Start a fresh page via the heading's own property instead of a manual
    # page break: idempotent, so it never leaves a trailing blank page even
    # when the previous content already ends exactly at a page boundary.
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(text)
    size = 16 if level == 1 else 14 if level == 2 else 13
    set_run_font(run, size=size, bold=True)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align="center", first_line=False, spacing_after=8)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, italic=True)


def set_cell_text(cell, text, bold=False, align="left", size=11):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, align=align, first_line=False, spacing_after=0)
    paragraph.paragraph_format.line_spacing = 1.15
    for index, line in enumerate(str(text).splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def keep_table_rows_together(table):
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if index == 0:
            header = tr_pr.find(qn("w:tblHeader"))
            if header is None:
                header = OxmlElement("w:tblHeader")
                tr_pr.append(header)
            header.set(qn("w:val"), "true")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align="center", size=11)
        set_cell_margins(table.rows[0].cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = "center" if idx == 0 or len(str(value)) <= 12 else "left"
            set_cell_text(cells[idx], value, align=align, size=11)
            set_cell_margins(cells[idx])
    if widths:
        set_table_width(table, widths)
    keep_table_rows_together(table)
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, align="left", first_line=False, spacing_after=4)
    return table


def add_figure_box(doc, text, caption):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align="center", first_line=False, spacing_after=8)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "8")
        element.set(qn("w:color"), "000000")
        p_bdr.append(element)
    p_pr.append(p_bdr)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=11, bold=False)
    add_caption(doc, caption)


def add_image_figure(doc, image_path, caption, width=6.4):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align="center", first_line=False, spacing_after=4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_toc_line(doc, label, page=None, indent=0):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, align="left", first_line=False, spacing_after=1)
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.25 * indent)
    if page is None:
        run = paragraph.add_run(label)
        set_run_font(run, size=13)
        return
    # Right-aligned tab stop with a dot leader so every page number lands
    # flush against the right text margin (6.5" = 8.5" page - 2x 1" margins).
    fmt.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=13)
    page_run = paragraph.add_run(f"\t{page}")
    set_run_font(page_run, size=13)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(0, 0, 0)


MERMAID_CONFIG = """{
  "theme": "default",
  "themeVariables": {
    "fontFamily": "Helvetica, Arial, sans-serif",
    "fontSize": "18px",
    "primaryColor": "#EEF0FF",
    "primaryBorderColor": "#6C63C7",
    "primaryTextColor": "#1b1b2a",
    "lineColor": "#4a4a68",
    "clusterBkg": "#f6f6ff",
    "clusterBorder": "#c9c9e6"
  },
  "flowchart": { "htmlLabels": true, "curve": "basis", "nodeSpacing": 55, "rankSpacing": 70, "padding": 12 },
  "er": { "layoutDirection": "TB", "entityPadding": 14 }
}
"""


# Each figure is authored as a Mermaid source so the diagrams stay consistent,
# editable and version-controlled. Rendered to PNG with the mermaid CLI (mmdc).
MERMAID_FIGURES = {
    "figure-1-architecture": """flowchart LR
  traffic["<b>Traffic Source</b><br/>demo user journeys<br/>API clients"]
  sut["<b>System Under Test</b><br/>backend/API<br/>structured logging"]
  logstore[("<b>Log Store</b><br/>structured API logs")]

  platform["<b>Behavioral Testing Platform</b><br/>log ingestion · behavior mining<br/>test generation · execution/review"]
  assist["<b>AI Assist</b><br/>naming, hints,<br/>triage"]
  storage[("<b>Record &amp; Artifact Storage</b><br/>decisions, history,<br/>specs, goldens, reports")]

  traffic -->|API traffic| sut
  sut -->|structured logs| logstore
  logstore -->|session/log data| platform
  platform -->|API replay| sut
  platform -->|records &amp; artifacts| storage
  storage -->|approved specs / goldens| platform
  assist -.->|support| platform
""",
    "figure-2-log-to-test": """flowchart LR
  log["<b>Structured log</b><br/>timestamp, method,<br/>endpoint, status"]
  clean["<b>Normalize &amp; clean</b><br/>mask id, drop noise,<br/>collapse retry"]
  session["<b>SessionFlow</b><br/>group by session_id,<br/>sort by time"]
  mine["<b>Mine candidates</b><br/>support, persona,<br/>priority"]
  plan["<b>Flow plan</b><br/>steps +<br/>expected status"]
  oracle["<b>Golden sync</b><br/>OpenAPI +<br/>observed schema"]
  spec["<b>Generated spec</b><br/>.spec.ts +<br/>fixtures"]
  approved["<b>HITL approval</b><br/>approved specs<br/>preserved"]
  log --> clean --> session --> mine --> plan --> oracle --> spec --> approved
""",
    "figure-3-golden-oracle": """flowchart LR
  live["<b>Live response</b><br/>status + JSON body"]
  oas["<b>OpenAPI schema</b><br/>authoritative contract"]
  observed["<b>Observed schema</b><br/>bodies-on log evidence"]
  rules["<b>Ignore / value rules</b><br/>dynamic IDs,<br/>timestamps, tokens"]
  compare["<b>compareResponse()</b><br/>schemaDiff + valueDiff"]
  result["<b>Deterministic pass / fail</b><br/>attach golden diff"]
  oas --> rules
  observed --> rules
  rules --> compare
  live --> compare
  compare --> result
""",
    "figure-4-regression-workflow": """flowchart LR
  mine["<b>Mine</b><br/>behavior<br/>candidates"]
  review["<b>Review</b><br/>dashboard<br/>HITL approve"]
  generate["<b>Generate</b><br/>Playwright<br/>API specs"]
  run["<b>Run</b><br/>test:all /<br/>persona target"]
  normalize["<b>Normalize</b><br/>Playwright JSON<br/>to normalized"]
  report["<b>Report</b><br/>green / red /<br/>invalid"]
  triage["<b>Triage sidecar</b><br/>advisory only"]
  persist["<b>Persist</b><br/>PostgreSQL +<br/>MinIO"]
  mine --> review --> generate --> run --> normalize --> report --> triage --> persist
""",
    "figure-5-postgres-erd": """erDiagram
  DECISIONS {
    uuid review_id PK
    text flow_signature
    text status
    jsonb payload
  }
  MANIFEST {
    uuid review_id PK
    jsonb payload
  }
  RUN_INDEX {
    text slug PK
    text status
    jsonb totals
  }
  INVARIANTS {
    uuid id PK
    text flow_signature
    bool verified
    jsonb payload
  }
  DISMISSED_RELATIONSHIPS {
    text pair_key PK
    jsonb payload
  }
  STORAGE_METADATA {
    text key PK
    jsonb payload
  }
  DECISIONS ||--o| MANIFEST : "review_id"
  DECISIONS ||--o{ RUN_INDEX : "run history"
  DECISIONS ||--o{ INVARIANTS : "flow_signature"
  INVARIANTS ||--o{ DISMISSED_RELATIONSHIPS : "flow pair"
  MANIFEST ||--o{ STORAGE_METADATA : "metadata"
""",
    "figure-6-minio-layout": """flowchart TB
  store[("<b>Kho đối tượng</b><br/>lưu trữ tương thích S3")]
  subgraph tests ["Sản phẩm kiểm thử"]
    direction LR
    aspecs["<b>Bộ kiểm thử<br/>đã phê duyệt</b>"]
    specs["<b>Bộ kiểm thử<br/>hiện hành</b>"]
    goldens["<b>Lược đồ<br/>tham chiếu</b>"]
    reports["<b>Báo cáo<br/>kết quả</b>"]
  end
  subgraph pipeline ["Dữ liệu quy trình"]
    direction LR
    sessions["<b>Luồng phiên<br/>người dùng</b>"]
    candidates["<b>Luồng hành vi<br/>ứng viên</b>"]
    validation["<b>Kết quả<br/>đánh giá</b>"]
    endpoint["<b>Ghi chú<br/>điểm cuối</b>"]
  end
  store --> tests
  store --> pipeline
""",
}


def render_architecture_figure(output_path):
    """Draw Figure 1 with fixed layout so the overview remains readable in DOCX."""
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    def load_font(size, bold=False):
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
            "/System/Library/Fonts/Supplemental/DejaVuSans-Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/DejaVuSans.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
        return ImageFont.load_default()

    font_title = load_font(30, bold=True)
    font_body = load_font(25)
    font_label = load_font(22)
    font_group = load_font(20)

    purple = "#8B6EEA"
    box_fill = "#EEF0FF"
    group_fill = "#F7F7FF"
    line = "#4A4A68"
    label_fill = "#ECECEC"

    def text_size(text, font):
        bbox = draw.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]

    def wrap_line(text, font, max_width):
        words = text.split()
        lines = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if text_size(candidate, font)[0] <= max_width or not current:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines

    def draw_centered_text(rect, title, body_lines=()):
        x1, y1, x2, y2 = rect
        lines = []
        for title_line in title.splitlines():
            lines.append((title_line, font_title))
        for body in body_lines:
            for line_text in wrap_line(body, font_body, x2 - x1 - 42):
                lines.append((line_text, font_body))
        line_heights = [text_size(text, font)[1] + 8 for text, font in lines]
        total_h = sum(line_heights) - 8
        y = y1 + ((y2 - y1) - total_h) / 2
        for (text, font), line_h in zip(lines, line_heights):
            tw, th = text_size(text, font)
            draw.text((x1 + ((x2 - x1) - tw) / 2, y), text, fill="#1B1B2A", font=font)
            y += line_h

    def draw_box(rect, title, body_lines=()):
        draw.rounded_rectangle(rect, radius=6, fill=box_fill, outline=purple, width=3)
        draw_centered_text(rect, title, body_lines)

    def draw_store(rect, title, body_lines=()):
        x1, y1, x2, y2 = rect
        ellipse_h = 34
        draw.rectangle((x1, y1 + ellipse_h / 2, x2, y2 - ellipse_h / 2), fill=box_fill, outline=purple, width=3)
        draw.ellipse((x1, y1, x2, y1 + ellipse_h), fill=box_fill, outline=purple, width=3)
        draw.arc((x1, y2 - ellipse_h, x2, y2), 0, 180, fill=purple, width=3)
        draw.line((x1, y1 + ellipse_h / 2, x1, y2 - ellipse_h / 2), fill=purple, width=3)
        draw.line((x2, y1 + ellipse_h / 2, x2, y2 - ellipse_h / 2), fill=purple, width=3)
        draw_centered_text(rect, title, body_lines)

    def arrow_head(p1, p2, fill=line):
        import math

        x1, y1 = p1
        x2, y2 = p2
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 13
        left = (x2 - size * math.cos(angle - 0.45), y2 - size * math.sin(angle - 0.45))
        right = (x2 - size * math.cos(angle + 0.45), y2 - size * math.sin(angle + 0.45))
        draw.polygon([p2, left, right], fill=fill)

    def draw_polyline(points, dashed=False):
        if dashed:
            for start, end in zip(points, points[1:]):
                draw_dashed_segment(start, end)
        else:
            draw.line(points, fill=line, width=3, joint="curve")
        arrow_head(points[-2], points[-1])

    def draw_dashed_segment(start, end, dash=16, gap=10):
        import math

        x1, y1 = start
        x2, y2 = end
        dx, dy = x2 - x1, y2 - y1
        dist = math.hypot(dx, dy)
        if dist == 0:
            return
        ux, uy = dx / dist, dy / dist
        pos = 0
        while pos < dist:
            seg_end = min(pos + dash, dist)
            draw.line((x1 + ux * pos, y1 + uy * pos, x1 + ux * seg_end, y1 + uy * seg_end), fill=line, width=3)
            pos += dash + gap

    def draw_label(text, xy):
        x, y = xy
        wrapped = wrap_line(text, font_label, 220)
        widths = [text_size(item, font_label)[0] for item in wrapped]
        line_h = text_size("Ag", font_label)[1] + 5
        w = max(widths) + 18
        h = len(wrapped) * line_h + 8
        draw.rounded_rectangle((x - w / 2, y - h / 2, x + w / 2, y + h / 2), radius=3, fill=label_fill)
        ty = y - h / 2 + 4
        for item in wrapped:
            tw, _ = text_size(item, font_label)
            draw.text((x - tw / 2, ty), item, fill="#252525", font=font_label)
            ty += line_h

    sources_rect = (70, 60, 1730, 265)
    platform_rect = (250, 355, 1550, 680)
    support_rect = (250, 785, 1550, 990)

    for rect, title in [
        (sources_rect, "Nguồn lưu lượng và hệ thống được kiểm thử"),
        (platform_rect, "Nền tảng kiểm thử hành vi"),
        (support_rect, "Điều phối, AI hỗ trợ và lưu trữ"),
    ]:
        draw.rounded_rectangle(rect, radius=8, fill=group_fill, outline="#C9C9E6", width=2)
        draw.text((rect[0] + 18, rect[1] + 8), title, fill="#333344", font=font_group)

    users = (120, 115, 430, 235)
    sut = (560, 115, 895, 235)
    logs = (1035, 105, 1375, 245)
    ingest = (310, 435, 570, 575)
    mine = (635, 435, 895, 575)
    generate = (960, 435, 1220, 575)
    run = (1285, 435, 1510, 575)
    dashboard = (320, 835, 610, 950)
    ai = (760, 835, 1040, 950)
    storage = (1195, 815, 1520, 965)

    draw_box(users, "API clients", ["traffic"])
    draw_box(sut, "Backend/API", ["under test", "structured log"])
    draw_store(logs, "Log store", ["structured logs"])
    draw_box(ingest, "Ingestion", ["clean", "sessionize"])
    draw_box(mine, "Behavior\nengine", ["flows", "persona"])
    draw_box(generate, "Generator", ["specs", "oracle"])
    draw_box(run, "Runner /\nreport", ["replay", "result"])
    draw_box(dashboard, "Dashboard", ["review", "approve/run"])
    draw_box(ai, "AI assist", ["naming", "triage"])
    draw_store(storage, "Storage", ["records", "artifacts"])

    draw_polyline([(430, 175), (560, 175)])
    draw_label("API traffic", (495, 145))
    draw_polyline([(895, 175), (1035, 175)])
    draw_label("structured logs", (965, 145))
    draw_polyline([(1205, 245), (1205, 315), (440, 315), (440, 435)])
    draw_label("session/log data", (815, 292))

    draw_polyline([(570, 505), (635, 505)])
    draw_polyline([(895, 505), (960, 505)])
    draw_polyline([(1220, 505), (1285, 505)])

    draw_polyline([(1395, 435), (1395, 315), (730, 315), (730, 235)])
    draw_label("API replay", (1085, 292))
    draw_polyline([(465, 835), (465, 720), (1090, 720), (1090, 575)])
    draw_polyline([(1395, 575), (1395, 720), (465, 720), (465, 835)])
    draw_label("result", (950, 742))
    draw_polyline([(900, 835), (900, 715), (765, 575)], dashed=True)
    draw_polyline([(900, 835), (900, 715), (1090, 575)], dashed=True)
    draw_polyline([(1395, 575), (1395, 815)])
    draw_label("artifacts", (1465, 700))
    draw_polyline([(1195, 890), (1090, 890), (1090, 575)])
    draw_label("approved data", (1090, 825))

    image.save(output_path)


def generate_report_figures():
    """Render every report figure from its Mermaid source with the mermaid CLI.

    Falls back with a clear error if `mmdc` is unavailable so the failure is
    obvious instead of silently shipping stale images.
    """
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    mermaid_dir = FIGURE_DIR / "mermaid"
    mermaid_dir.mkdir(parents=True, exist_ok=True)

    mmdc = shutil.which("mmdc")
    if not mmdc:
        raise RuntimeError(
            "mermaid CLI (mmdc) not found on PATH. Install with "
            "`npm install -g @mermaid-js/mermaid-cli` to regenerate figures."
        )

    config_path = mermaid_dir / "config.json"
    config_path.write_text(MERMAID_CONFIG, encoding="utf-8")

    for name, source in MERMAID_FIGURES.items():
        mmd_path = mermaid_dir / f"{name}.mmd"
        png_path = FIGURE_DIR / f"{name}.png"
        mmd_path.write_text(source, encoding="utf-8")
        subprocess.run(
            [
                mmdc,
                "-i", str(mmd_path),
                "-o", str(png_path),
                "-c", str(config_path),
                "-b", "white",
                "--width", "1700",
            ],
            check=True,
        )


def add_cover(doc):
    for line in [
        "TẬP ĐOÀN CÔNG NGHIỆP - VIỄN THÔNG QUÂN ĐỘI",
        "VIETTEL DIGITAL TALENT 2026",
    ]:
        add_paragraph(doc, line, bold=True, align="center", size=14)
    for _ in range(3):
        doc.add_paragraph()
    add_paragraph(doc, REPORT_TITLE, bold=True, align="center", size=18)
    add_paragraph(doc, TOPIC, bold=True, italic=True, align="center", size=16)
    doc.add_paragraph()
    add_paragraph(doc, AUTHOR, bold=True, align="center", size=14)
    add_paragraph(doc, EMAIL, align="center", size=13)
    add_paragraph(doc, f"Chương trình: {PROGRAM}", align="center", size=13)
    add_paragraph(doc, f"Lĩnh vực: {FIELD}", align="center", size=13)
    for _ in range(3):
        doc.add_paragraph()
    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [["Mentor", MENTOR], ["Đơn vị", UNIT]],
        widths=[1.6, 4.8],
    )
    add_paragraph(doc, "Hà Nội, 2026", align="center", size=13)


def add_front_matter(doc):
    add_heading(doc, "LỜI MỞ ĐẦU", 1, page_break_before=True)
    for text in [
        "Mini-project này xuất phát từ nhu cầu giảm công sức kiểm thử hồi quy cho backend/API. Trong một hệ thống thương mại điện tử, một thay đổi nhỏ ở các nghiệp vụ như giỏ hàng, thanh toán, đơn hàng hay trả hàng có thể làm hỏng nhiều luồng nghiệp vụ vốn đã chạy ổn định. Nếu chỉ viết kiểm thử thủ công theo tài liệu yêu cầu, bộ kiểm thử dễ bỏ sót các chuỗi thao tác thực sự diễn ra khi người dùng tương tác với hệ thống.",
        "Ý tưởng cốt lõi của đề tài là tận dụng nhật ký (log) mà backend sinh ra trong quá trình vận hành làm nguồn dữ liệu để tự động dựng lại các kịch bản kiểm thử. Hệ thống thu thập log có cấu trúc, phân tích để tái hiện hành vi người dùng, rồi sinh ra các bài kiểm thử API tương ứng và chạy lại chúng để phát hiện hồi quy. Một hệ thống thương mại điện tử mã nguồn mở được sử dụng làm đối tượng kiểm thử (System Under Test) cho toàn bộ quá trình minh họa.",
        "Báo cáo tập trung chứng minh rằng quy trình này chạy được trọn vẹn trong môi trường thử nghiệm, chứ không dừng ở mức ý tưởng. Dữ liệu đánh giá hiện tại đến từ lưu lượng mô phỏng và hạ tầng lưu trữ cục bộ; báo cáo không khẳng định đã xử lý lưu lượng vận hành thực tế.",
        "Trí tuệ nhân tạo (AI/LLM) trong hệ thống được đặt ở vai trò hỗ trợ: tạo lưu lượng đa dạng, đặt tên luồng, gợi ý điểm kiểm chứng hoặc phân loại kết quả. Việc quyết định một bài kiểm thử đạt hay không đạt vẫn dựa trên các cơ chế tất định (deterministic) để kết quả có thể kiểm toán và tái lập.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "TÓM TẮT NỘI DUNG VÀ ĐÓNG GÓP", 1, page_break_before=True)
    for text in [
        "Đề tài xây dựng một quy trình khép kín biến nhật ký backend thành các bài kiểm thử hồi quy cho API. Hệ thống được tổ chức thành các thành phần chính có vai trò rõ ràng: sinh lưu lượng mô phỏng, ghi và thu thập log, chuẩn hóa log thành luồng phiên người dùng, khai phá hành vi, sinh và chạy kiểm thử, tổng hợp báo cáo, cùng giao diện quản trị và lớp lưu trữ.",
        "Kết quả thực tế của hệ thống là một chuỗi sản phẩm có thể quan sát được: luồng phiên người dùng, các luồng hành vi ứng viên, bộ kiểm thử API được sinh tự động, lược đồ phản hồi tham chiếu, kết quả chạy đã chuẩn hóa, báo cáo và các chỉ số đánh giá. Bản chụp dữ liệu ngày 07/07/2026 ghi nhận 18 luồng ứng viên, 9 bộ kiểm thử đã phê duyệt, 56 golden schema theo cặp endpoint/trạng thái phản hồi, 295 bản ghi invariant trong PostgreSQL (113 đã xác minh, 182 chưa xác minh) và lượt regression mới nhất chạy xanh với 9/9 kiểm thử đạt.",
        "Đóng góp của mini-project nằm ở việc biến log backend thành các bài kiểm thử hồi quy có nguồn gốc truy vết được, kết hợp phương pháp khai phá chuỗi tất định với cơ chế tham chiếu dựa trên đặc tả API. Đề tài cũng phân định rõ ranh giới của phần thử nghiệm: lưu lượng hiện tại là mô phỏng, hạ tầng chạy cục bộ, đánh giá đột biến mới nhất đạt 40/45 mutant đo được, và AI không nằm trên đường quyết định kết quả kiểm thử.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "MỤC LỤC", 1, page_break_before=True)
    toc_items = [
        ("Lời mở đầu", "2", 0),
        ("Tóm tắt nội dung và đóng góp", "3", 0),
        ("Danh mục hình vẽ", "5", 0),
        ("Danh mục bảng biểu", "5", 0),
        ("I. Giới thiệu", "6", 0),
        ("1.1. Bối cảnh", "6", 1),
        ("1.2. Vấn đề đặt ra", "6", 1),
        ("1.3. Mục tiêu", "7", 1),
        ("1.4. Phạm vi", "7", 1),
        ("II. Nội dung và phương pháp", "8", 0),
        ("2.1. Tổng quan hệ thống hiện tại", "8", 1),
        ("2.2. Luồng xử lý chính", "9", 1),
        ("2.3. Log format và masking", "10", 1),
        ("2.4. Behavior engine", "11", 1),
        ("2.5. Test generator và golden oracle", "11", 1),
        ("2.6. Storage, dashboard và vai trò AI", "12", 1),
        ("III. Kết quả thực hiện và đánh giá", "16", 0),
        ("3.1. Môi trường thử nghiệm", "16", 1),
        ("3.2. Các flow được tái dựng và test đã sinh", "16", 1),
        ("3.3. Kết quả validation và report", "18", 1),
        ("3.4. Đánh giá", "19", 1),
        ("IV. Kết luận", "22", 0),
        ("Tài liệu tham khảo", "24", 0),
    ]
    for label, page, indent in toc_items:
        add_toc_line(doc, label, page, indent)

    add_heading(doc, "DANH MỤC HÌNH VẼ", 1, page_break_before=True)
    for item in [
        "Hình 1. Kiến trúc tổng quan pipeline log-to-test",
        "Hình 2. Luồng xử lý từ structured log đến generated API test",
        "Hình 3. Cơ chế golden oracle khi replay test",
        "Hình 4. Quy trình chạy regression test trong môi trường demo/CI",
        "Hình 5. Mô hình dữ liệu của lớp lưu trữ bản ghi",
        "Hình 6. Tổ chức sản phẩm trong kho lưu trữ đối tượng",
    ]:
        add_toc_line(doc, item)

    add_heading(doc, "DANH MỤC BẢNG BIỂU", 1)
    for item in [
        "Bảng 1. Các thành phần chính của hệ thống và vai trò",
        "Bảng 2. Các nhóm trường log/API sử dụng trong hệ thống",
        "Bảng 3. So sánh replay log, LLM-only và hybrid deterministic approach",
        "Bảng 4. Các flow được tái dựng từ log demo",
        "Bảng 5. Kết quả sinh/chạy test trong môi trường thử nghiệm",
        "Bảng 6. Hạn chế hiện tại và hướng cải thiện",
    ]:
        add_toc_line(doc, item)


def add_intro_chapter(doc):
    add_heading(doc, "I. GIỚI THIỆU", 1, page_break_before=True)
    add_heading(doc, "1.1. Bối cảnh", 2)
    for text in [
        "Kiểm thử hồi quy cho backend/API là bài toán trọng tâm của đề tài. Trong một hệ thống thương mại điện tử, các nghiệp vụ như đăng nhập, duyệt sản phẩm, giỏ hàng, thanh toán, đơn hàng, giao hàng và trả hàng đều trải qua nhiều bước gọi API. Khi backend thay đổi, kiểm thử hồi quy cần bảo đảm các luồng này vẫn hoạt động đúng mà không phụ thuộc vào thao tác trên giao diện.",
        "Nhật ký có cấu trúc phản ánh trung thực cách các yêu cầu thực sự đi qua hệ thống, bao gồm thao tác, điểm cuối được gọi, trạng thái phản hồi, thời điểm và định danh phiên. Nếu chỉ dùng log để gỡ lỗi, đội phát triển đã bỏ lỡ một nguồn dữ liệu có thể tái sử dụng trực tiếp cho kiểm thử.",
        "Ý tưởng của mini-project là tái dựng các luồng hành vi từ log và biến chúng thành các bài kiểm thử hồi quy ở mức API. Cách làm này giúp ưu tiên những chuỗi thao tác đã thực sự xuất hiện trong lưu lượng, thay vì chỉ dựa trên một danh sách điểm cuối rời rạc.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.2. Vấn đề đặt ra", 2)
    for text in [
        "Kiểm thử viết thủ công dễ bỏ sót các luồng thực tế như thanh toán sau khi đăng nhập, giỏ hàng thất bại do dữ liệu không hợp lệ, quản trị viên hủy đơn hay toàn bộ vòng đời trả hàng. Mỗi khi API thay đổi, việc cập nhật trạng thái kỳ vọng, các định danh động và cấu trúc phản hồi cũng tiêu tốn nhiều công sức.",
        "Việc phát lại các yêu cầu từ log không thể thực hiện một cách máy móc. Dữ liệu nhạy cảm phải được che, các điểm cuối chứa định danh động cần được chuẩn hóa, các yêu cầu lặp do thử lại hoặc bộ đệm phải được loại nhiễu, ngữ cảnh phiên và xác thực phải được tái tạo đúng, và các trường thay đổi theo thời gian không được đem ra so sánh cứng.",
        "Do đó, hệ thống cần một thành phần khai phá hành vi để chắt lọc những luồng có ý nghĩa, và một cơ chế tham chiếu để xác định hành vi kỳ vọng khi phát lại. Nếu chỉ dùng AI để sinh kiểm thử, kết quả có nguy cơ không tái lập được; ngược lại, nếu chỉ phát lại log một cách thô, tỷ lệ báo lỗi giả sẽ cao vì phản hồi thực luôn chứa các trường biến động.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.3. Mục tiêu", 2)
    for text in [
        "Mục tiêu cụ thể của đề tài là thu thập nhật ký API có cấu trúc, chuẩn hóa thành các luồng phiên người dùng, khai phá các luồng hành vi ứng viên, sinh bộ kiểm thử API tương ứng, chạy lại chúng và xuất báo cáo kết quả với các trạng thái đạt/không đạt/không hợp lệ.",
        "Mục tiêu kỹ thuật là giữ cho việc đánh giá kết quả mang tính tất định, có thể kiểm toán và tái lập. AI có thể hỗ trợ đặt tên, tạo lưu lượng, gợi ý điểm kiểm chứng hoặc phân loại kết quả, nhưng không quyết định một bài kiểm thử đạt hay không đạt.",
        "Mục tiêu đánh giá là chứng minh toàn bộ quy trình vận hành được trong môi trường thử nghiệm với lưu lượng mô phỏng và hạ tầng lưu trữ cục bộ, trong đó cơ sở dữ liệu quan hệ lưu các bản ghi và quyết định, còn kho đối tượng lưu các sản phẩm của quy trình.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.4. Phạm vi", 2)
    for text in [
        "Đề tài tập trung vào kiểm thử hồi quy ở mức API/backend, không bao gồm kiểm thử giao diện hay kiểm thử đầu-cuối trên trình duyệt. Bộ kiểm thử được sinh ra chỉ tương tác với hệ thống thông qua các lời gọi API.",
        "Một hệ thống thương mại điện tử mã nguồn mở được chọn làm đối tượng kiểm thử minh họa. Các luồng nghiệp vụ liên quan bao gồm đăng nhập, duyệt sản phẩm, giỏ hàng, thanh toán, đơn hàng, quản trị danh mục, giao hàng, hủy đơn và trả hàng.",
        "Dữ liệu hiện tại là lưu lượng mô phỏng trên hạ tầng cục bộ. Hệ thu thập log cùng cơ sở dữ liệu và kho lưu trữ đều chạy ở quy mô một nút, phù hợp cho mục đích thử nghiệm; báo cáo không khẳng định hệ thống đã kết nối lưu lượng vận hành thực tế hay sẵn sàng triển khai sản phẩm.",
    ]:
        add_paragraph(doc, text)


def add_method_chapter(doc):
    add_heading(doc, "II. NỘI DUNG VÀ PHƯƠNG PHÁP", 1, page_break_before=True)
    add_heading(doc, "2.1. Tổng quan hệ thống hiện tại", 2)
    add_paragraph(
        doc,
        "Hệ thống được tổ chức thành nhiều thành phần độc lập nhưng phối hợp theo một quy trình xuyên suốt, đi từ lưu lượng mô phỏng, qua thu thập và chuẩn hóa log, khai phá hành vi, sinh và chạy kiểm thử, cho tới tổng hợp báo cáo. Toàn bộ được vận hành trên nền tảng chứa (container) để dễ dựng lại môi trường, với hệ thống được kiểm thử, giao diện quản trị, cơ sở dữ liệu, kho lưu trữ và hệ thu thập log cùng hoạt động. Bảng 1 tóm tắt vai trò của các thành phần chính; kiến trúc tổng quan được minh họa ở Hình 1.",
    )
    add_table(
        doc,
        ["Thành phần", "Vai trò trong hệ thống"],
        [
            ["Hệ thống được kiểm thử", "Backend thương mại điện tử đóng vai trò đối tượng kiểm thử; đồng thời ghi lại nhật ký có cấu trúc cho mỗi yêu cầu API."],
            ["Sinh lưu lượng", "Tạo lưu lượng mô phỏng cho nhiều kịch bản: duyệt hàng, thanh toán, trả hàng, thao tác quản trị và các trường hợp biên."],
            ["Thu thập log", "Thu gom, vận chuyển và lưu trữ nhật ký để phục vụ phân tích về sau."],
            ["Chuẩn hóa log", "Làm sạch, loại nhiễu và nhóm các yêu cầu thành các luồng phiên người dùng."],
            ["Khai phá hành vi", "Phát hiện các chuỗi thao tác lặp lại, suy luận nhóm người dùng và chắt lọc các luồng ứng viên."],
            ["Cơ chế tham chiếu", "Xác định hành vi kỳ vọng khi phát lại, dựa trên đặc tả API và bằng chứng quan sát được từ log."],
            ["Sinh & chạy kiểm thử", "Chuyển các luồng hành vi thành bộ kiểm thử API và thực thi lại chúng theo từng nhóm người dùng."],
            ["Tổng hợp báo cáo", "Chuẩn hóa kết quả chạy và dựng báo cáo với trạng thái đạt/không đạt cùng thông tin phân loại."],
            ["Giao diện & lưu trữ", "Cho phép rà soát, phê duyệt và điều phối quy trình; lưu trữ các bản ghi và sản phẩm của quy trình."],
        ],
        widths=[1.85, 4.6],
        caption="Bảng 1. Các thành phần chính của hệ thống và vai trò",
    )
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-1-architecture.png",
        "Hình 1. Kiến trúc tổng quan pipeline log-to-test",
        width=6.4,
    )

    add_heading(doc, "2.2. Luồng xử lý chính", 2)
    for text in [
        "Quy trình chuẩn gồm các bước tuần tự: khởi động hệ thống được kiểm thử và hạ tầng phụ trợ, sinh lưu lượng mô phỏng, chuyển log thô thành các luồng phiên người dùng, khai phá các luồng hành vi ứng viên, sinh bộ kiểm thử cùng dữ liệu tham chiếu, và cuối cùng là chạy lại toàn bộ kiểm thử. Mỗi bước nhận đầu ra của bước trước làm đầu vào, nhờ đó có thể chạy lại độc lập khi cần.",
        "Việc phân loại nhóm người dùng được suy luận từ chính đặc điểm của luồng thao tác, chẳng hạn dấu hiệu truy cập các chức năng cần đăng nhập hay các điểm cuối dành cho quản trị, chứ không lấy trực tiếp từ nhãn có sẵn. Nhãn vai trò quan sát được trong log chỉ được dùng làm dữ liệu đối chiếu để đánh giá độ chính xác của bước phân loại.",
        "Bộ kiểm thử được sinh ra được sắp xếp theo nhóm người dùng và tách thành hai loại: luồng thành công và luồng thất bại. Nhờ đó, các hành vi lỗi hợp lệ, như thao tác bị từ chối do dữ liệu không đúng, cũng được giữ lại thành kiểm thử riêng thay vì bị bỏ qua.",
    ]:
        add_paragraph(doc, text)
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-2-log-to-test.png",
        "Hình 2. Luồng xử lý từ structured log đến generated API test",
    )

    add_heading(doc, "2.3. Log format và masking", 2)
    add_paragraph(
        doc,
        "Mỗi yêu cầu API được ghi lại dưới dạng một bản ghi có cấu trúc, bao gồm các trường mô tả thời điểm, phiên làm việc, thao tác, điểm cuối, trạng thái phản hồi và ngữ cảnh xác thực. Nội dung chi tiết của thân yêu cầu và phản hồi chỉ được ghi khi bật tùy chọn tương ứng, và luôn đi qua bước che dữ liệu nhạy cảm trước khi lưu. Bảng 2 tóm tắt vai trò của các nhóm trường chính đối với quy trình.",
    )
    add_table(
        doc,
        ["Nhóm trường", "Vai trò trong quy trình"],
        [
            ["Thời điểm", "Sắp xếp thứ tự các yêu cầu trong một luồng phiên."],
            ["Định danh phiên", "Nhóm các yêu cầu rời rạc thành một hành vi người dùng."],
            ["Định danh truy vết", "Theo dõi hành trình của yêu cầu; không đưa vào so sánh cứng."],
            ["Thao tác & điểm cuối", "Tạo biểu diễn chuẩn hóa cho từng bước gọi API."],
            ["Trạng thái phản hồi", "Xác định trạng thái kỳ vọng và nhận diện luồng lỗi."],
            ["Ngữ cảnh người dùng", "Dùng làm dữ liệu đối chiếu để đánh giá phân loại."],
            ["Thân yêu cầu / phản hồi", "Tạo lược đồ tham chiếu; đã được rút gọn và che dữ liệu."],
        ],
        widths=[2.1, 4.35],
        caption="Bảng 2. Các nhóm trường log/API sử dụng trong hệ thống",
    )
    add_paragraph(
        doc,
        "Cơ chế che dữ liệu tự động nhận diện và thay thế các trường nhạy cảm như mật khẩu, mã xác thực, thông tin thanh toán, liên hệ cá nhân và địa chỉ. Điểm đáng chú ý là cấu trúc (hình dạng) của dữ liệu vẫn được giữ nguyên trong khi giá trị thật bị thay bằng giá trị an toàn, nhờ đó lược đồ tham chiếu vẫn có đủ bằng chứng để đối chiếu mà không làm lộ thông tin.",
    )

    add_heading(doc, "2.4. Behavior engine", 2)
    for text in [
        "Trước khi khai phá, dữ liệu log được chuẩn hóa: các điểm cuối chứa định danh động được đưa về dạng khái quát, các yêu cầu nhiễu như kiểm tra sức khỏe hệ thống, tài nguyên tĩnh hay phản hồi từ bộ đệm bị loại bỏ, và các yêu cầu lặp do thử lại được gộp lại. Sau đó, các yêu cầu được nhóm theo phiên và sắp xếp theo thời gian để tạo thành luồng hành vi.",
        "Thành phần khai phá hành vi biểu diễn mỗi bước dưới dạng một biểu tượng chuẩn hóa, rồi áp dụng các thuật toán khai phá chuỗi tuần tự để tìm các mẫu thao tác lặp lại. Các luồng thu được được khử trùng lặp, xếp hạng theo mức độ ưu tiên và giới hạn số lượng theo từng nhóm người dùng. Trạng thái kỳ vọng của mỗi bước được xác định theo giá trị phổ biến nhất trong cùng một ngữ cảnh xác thực, nhằm tránh trộn lẫn thất bại của khách vãng lai với thành công của khách đã đăng nhập.",
        "Việc khai phá luồng lỗi được tách riêng: với những phiên xuất hiện lỗi, hệ thống lấy đoạn thao tác dẫn tới điểm lỗi và chỉ giữ lại khi mẫu lỗi đó xuất hiện đủ nhiều. Nhờ vậy, các hành vi tiêu cực, như thao tác giỏ hàng bị từ chối do dữ liệu không hợp lệ, được giữ thành kiểm thử luồng thất bại thay vì bị che khuất bởi các luồng thành công phổ biến hơn.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.5. Test generator và golden oracle", 2)
    for text in [
        "Thành phần sinh kiểm thử chuyển mỗi luồng hành vi thành một bài kiểm thử API cùng dữ liệu phụ trợ đi kèm. Mỗi bài kiểm thử được gắn một chữ ký nhận dạng để hệ thống biết luồng nào đã được phê duyệt và cần bảo toàn qua các lần sinh lại. Với luồng thành công, hệ thống bắt buộc phải dựng được dữ liệu tham chiếu trước khi phát hành kiểm thử; nếu thiếu, bài kiểm thử bị từ chối thay vì suy biến thành một kiểm thử chỉ so trạng thái.",
        "Cơ chế tham chiếu kết hợp đặc tả API chính thức với lược đồ quan sát được từ log, cùng một tập quy tắc bỏ qua và quy tắc giá trị. Khi phát lại, hệ thống tra cứu dữ liệu tham chiếu theo điểm cuối và trạng thái, đối chiếu sự khác biệt về cấu trúc và giá trị, rồi đính kèm kết quả so sánh vào thông tin kiểm thử. Các trường biến động không bị đem ra so sánh cứng, nhằm giảm tỷ lệ báo lỗi giả khi phát lại.",
    ]:
        add_paragraph(doc, text)
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-3-golden-oracle.png",
        "Hình 3. Cơ chế golden oracle khi replay test",
    )
    add_table(
        doc,
        ["Cách tiếp cận", "Ưu điểm", "Rủi ro", "Cách hệ thống xử lý"],
        [
            ["Phát lại log thô", "Nhanh, bám sát yêu cầu thật.", "Dễ báo lỗi giả vì các định danh, thời điểm và tác động phụ luôn thay đổi.", "Không phát lại thô; chuẩn hóa luồng và dùng cơ chế tham chiếu."],
            ["Chỉ dùng AI sinh kiểm thử", "Mô tả tự nhiên, gợi ý nhiều trường hợp.", "Có thể bịa nội dung, điểm kiểm chứng khó kiểm toán.", "AI chỉ hỗ trợ đặt tên/gợi ý/phân loại, không quyết định kết quả."],
            ["Kết hợp tất định (hướng của đề tài)", "Có nguồn gốc từ log, kết quả tái lập, phù hợp tích hợp liên tục.", "Quy trình phức tạp hơn, cần nhiều lớp kiểm soát.", "Là hướng tiếp cận chính của hệ thống."],
        ],
        widths=[1.35, 1.65, 1.65, 1.85],
        caption="Bảng 3. So sánh replay log, LLM-only và hybrid deterministic approach",
    )

    add_heading(doc, "2.6. Storage, dashboard và vai trò AI", 2)
    for text in [
        "Lớp lưu trữ được thiết kế trừu tượng và tách thành hai loại: lưu trữ bản ghi có cấu trúc và lưu trữ đối tượng nhị phân. Ở cấu hình cục bộ, dữ liệu có thể nằm trên hệ thống tệp; ở cấu hình từ xa, các bản ghi và ràng buộc nghiệp vụ được đặt trong cơ sở dữ liệu quan hệ, còn các sản phẩm kích thước lớn được đặt trong kho đối tượng. Nhờ tách bạch như vậy, giao diện quản trị và quy trình đều thao tác qua cùng một lớp trừu tượng bất kể dữ liệu nằm ở đâu.",
        "Trong cấu hình từ xa, cơ sở dữ liệu quan hệ lưu các quyết định phê duyệt, lịch sử chạy và các ràng buộc nghiệp vụ, trong khi kho đối tượng lưu các bộ kiểm thử, luồng ứng viên, lược đồ tham chiếu, báo cáo và luồng phiên. Cách tổ chức này cho thấy hệ thống không phụ thuộc vào các tệp cục bộ trong mã nguồn mà đã tách dữ liệu ra một lớp lưu trữ độc lập. Mô hình dữ liệu và cách tổ chức sản phẩm được minh họa ở Hình 5 và Hình 6.",
        "AI xuất hiện ở hai vị trí: hỗ trợ tạo lưu lượng đa dạng và hỗ trợ đặt tên hoặc sửa chữa trong khâu sinh kiểm thử. Đây đều là các tác vụ làm giàu dữ liệu. Toàn bộ các khâu mang tính quyết định, gồm phân loại người dùng, đối chiếu tham chiếu, đánh giá đạt/không đạt và trạng thái báo cáo, vẫn được thực hiện theo quy tắc tất định.",
    ]:
        add_paragraph(doc, text)
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-4-regression-workflow.png",
        "Hình 4. Quy trình chạy regression test trong môi trường demo/CI",
    )
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-5-postgres-erd.png",
        "Hình 5. Mô hình dữ liệu của lớp lưu trữ bản ghi",
    )
    add_image_figure(
        doc,
        FIGURE_DIR / "figure-6-minio-layout.png",
        "Hình 6. Tổ chức sản phẩm trong kho lưu trữ đối tượng",
    )


def add_results_chapter(doc):
    add_heading(doc, "III. KẾT QUẢ THỰC HIỆN VÀ ĐÁNH GIÁ", 1, page_break_before=True)
    add_heading(doc, "3.1. Môi trường thử nghiệm", 2)
    for text in [
        "Môi trường thử nghiệm được dựng hoàn toàn bằng nền tảng chứa (container) trên máy phát triển. Toàn bộ các thành phần chính đều hoạt động đồng thời, bao gồm hệ thống được kiểm thử, cơ sở dữ liệu, kho lưu trữ đối tượng, hệ thu thập log và giao diện quản trị.",
        "Dữ liệu trình bày trong chương này được lấy trực tiếp từ cơ sở dữ liệu và kho lưu trữ của hệ thống. Đây là dữ liệu sinh ra từ lưu lượng mô phỏng, không phải lưu lượng vận hành thực tế; hạ tầng chạy ở quy mô một nút, phù hợp cho thử nghiệm nhưng chưa thể xem là cấu hình triển khai sản phẩm.",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        ["Nguồn", "Số liệu ghi nhận", "Ý nghĩa"],
        [
            ["Cơ sở dữ liệu bản ghi", "9 quyết định đã phê duyệt, 11 lượt chạy, 10 bản kê, 295 bản ghi invariant (113 đã xác minh, 182 chưa xác minh)", "Kết quả rà soát, lịch sử chạy và siêu dữ liệu kiểm chứng đã được lưu bền vững."],
            ["Kho lưu trữ đối tượng", "9 bộ kiểm thử đã phê duyệt, 56 golden schema theo endpoint/trạng thái, 33 báo cáo, 29 bộ kiểm thử hiện hành", "Các sản phẩm của quy trình được lưu trên kho đối tượng độc lập."],
            ["Báo cáo chạy mới nhất", "9 kiểm thử được chạy, 9 đạt, 0 không đạt", "Lượt chạy hồi quy gần nhất đạt trạng thái xanh trong môi trường thử nghiệm."],
            ["Kết quả đánh giá", "Mutation score 88,9% trên 45 mutant đo được; 40 killed, 5 survived, 105 inconclusive", "Tín hiệu hồi quy tốt trên mutant đo được; cần tăng độ phủ cho nhóm inconclusive."],
        ],
        widths=[1.55, 2.55, 2.35],
        caption="Bảng phụ. Tổng hợp số liệu từ cơ sở dữ liệu, kho lưu trữ và báo cáo",
    )

    add_heading(doc, "3.2. Các flow được tái dựng và test đã sinh", 2)
    add_paragraph(
        doc,
        "Ở lần khai phá gần nhất (run 2026-07-07T01-30-53-137Z), hệ thống chắt lọc được 18 luồng hành vi ứng viên từ tập log mô phỏng, với ngưỡng hỗ trợ tối thiểu bằng ba. Các luồng này được phân bổ thành 3 luồng khách vãng lai, 10 luồng khách đã đăng nhập và 5 luồng quản trị viên; PrefixSpan phát hiện 127.764 mẫu so với 189 mẫu n-gram baseline. Bảng 4 liệt kê các luồng tiêu biểu đã được tái dựng và chuyển thành kiểm thử.",
    )
    add_table(
        doc,
        ["Luồng nghiệp vụ", "Nhóm người dùng", "Loại luồng", "Kết quả"],
        [
            ["Duyệt danh sách sản phẩm", "Khách vãng lai", "Thành công", "Đạt"],
            ["Xem chi tiết sản phẩm", "Khách vãng lai", "Thành công", "Đạt"],
            ["Thanh toán có áp khuyến mãi", "Khách đã đăng nhập", "Thành công", "Đạt"],
            ["Thêm sản phẩm vào giỏ thất bại", "Khách đã đăng nhập", "Thất bại", "Đạt"],
            ["Cập nhật giỏ hàng thất bại", "Khách đã đăng nhập", "Thất bại", "Đạt"],
            ["Tạo phiếu giao hàng cho đơn", "Quản trị viên", "Thành công", "Đạt"],
            ["Hủy đơn hàng", "Quản trị viên", "Thành công", "Đạt"],
            ["Vòng đời trả hàng đầy đủ", "Quản trị viên", "Thành công", "Đạt"],
            ["Cập nhật sản phẩm (quản trị)", "Quản trị viên", "Thành công", "Đạt"],
        ],
        widths=[2.25, 1.65, 1.15, 0.9],
        caption="Bảng 4. Các luồng nghiệp vụ được tái dựng từ log thử nghiệm",
    )

    doc.add_page_break()

    add_table(
        doc,
        ["Nhóm kiểm thử", "Số bài", "Đạt", "Không đạt", "Ghi chú"],
        [
            ["Khách vãng lai", "2", "2", "0", "Duyệt sản phẩm và xem chi tiết sản phẩm."],
            ["Khách đã đăng nhập", "3", "3", "0", "Thanh toán khuyến mãi và hai luồng giỏ hàng thất bại."],
            ["Quản trị viên", "4", "4", "0", "Cập nhật sản phẩm, giao hàng, hủy đơn và trả hàng."],
            ["Tổng lượt chạy gần nhất", "9", "9", "0", "Toàn bộ đạt, trạng thái xanh."],
        ],
        widths=[1.75, 0.75, 0.65, 0.9, 2.5],
        caption="Bảng 5. Kết quả sinh và chạy kiểm thử trong môi trường thử nghiệm",
    )

    add_heading(doc, "3.3. Kết quả validation và report", 2)
    for text in [
        "Báo cáo validation mới nhất chấm 331 phiên mô phỏng. Luồng holdout của khách đã đăng nhập (đăng ký, tạo giỏ, thêm sản phẩm, thanh toán) được tái dựng với support 6/6, đúng ngưỡng đặt ra. Negative control cũng đạt: hệ thống không tạo ra luồng bất thường nằm ngoài dữ liệu đầu vào.",
        "Ở bước phân loại nhóm người dùng, biến thể cart-read signal đạt accuracy 0,9547 và macro-F1 0,9646 trên tập phiên có nhãn đối chiếu. Đây là kết quả đánh giá nội bộ trên dữ liệu thử nghiệm, không nên diễn giải thành chỉ số chuẩn cho môi trường vận hành.",
        "Báo cáo của lượt chạy gần nhất (run-2026-07-07-104122) đạt trạng thái xanh với 9 kiểm thử được chạy và đều đạt. Đáng chú ý, lượt run-2026-07-07-102319 trước đó từng ở trạng thái đỏ với 8 kiểm thử đạt và 1 kiểm thử không đạt, cho thấy cơ chế báo cáo có khả năng phản ánh hồi quy thực sự chứ không phải luôn báo xanh.",
        "Đánh giá đột biến mới nhất chạy trên toàn bộ mục tiêu với baseline sạch và executability 100%. Trong 150 mutant được sinh, 40 mutant bị phát hiện, 5 mutant sống sót và 105 mutant inconclusive; mutation score đạt 88,9% khi tính trên 45 mutant đo được. Các mutant bị phát hiện chủ yếu thuộc nhóm vi phạm enum và format, còn các mutant inconclusive cho thấy cần tăng thêm luồng baseline để exercise đều hơn các endpoint.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.4. Đánh giá", 2)
    for text in [
        "Những ưu điểm đã được chứng minh qua thực nghiệm gồm: quy trình khép kín với các thành phần phân tách rõ ràng; các kiểm thử được sinh ra chạy trực tiếp ở mức API; cơ chế tham chiếu từ chối phát hành khi thiếu dữ liệu đối chiếu thay vì bỏ qua; các bộ kiểm thử đã phê duyệt được bảo toàn; báo cáo truy vết được kết quả theo nhóm người dùng và luồng nghiệp vụ; và lớp lưu trữ tách bạch giữa bản ghi và sản phẩm.",
        "Bên cạnh các kết quả đã chạy được, cần nhìn nhận hệ thống ở nhiều khía cạnh rộng hơn. Nguồn dữ liệu hiện tại vẫn là lưu lượng mô phỏng, vì vậy chưa phản ánh đầy đủ độ nhiễu của môi trường vận hành thực tế như phiên bị gián đoạn, retry, timeout, cache, request bất đồng bộ, race condition hoặc hành vi trải dài qua nhiều dịch vụ. Khi chuyển sang log thật, việc tái dựng luồng có thể khó hơn đáng kể so với môi trường cục bộ.",
        "Một hạn chế khác nằm ở chất lượng log đầu vào. Pipeline hiện giả định backend ghi nhận đủ thông tin về phiên, endpoint, trạng thái phản hồi, ngữ cảnh xác thực và một phần request/response đã được che dữ liệu. Trong thực tế, log giữa các service có thể không đồng nhất, thiếu correlation id, thiếu ngữ cảnh nghiệp vụ hoặc bị rút gọn vì lý do chi phí và bảo mật. Nếu log không đủ giàu thông tin, hệ thống chỉ có thể sinh test ở mức bề mặt thay vì tái hiện chính xác ý nghĩa nghiệp vụ.",
        "Cơ chế golden oracle hiện tại giúp giảm báo lỗi giả bằng cách kết hợp đặc tả API, lược đồ quan sát được và các quy tắc bỏ qua trường động, nhưng vẫn chưa đủ để khẳng định toàn bộ hành vi nghiệp vụ là đúng. Nhiều lỗi backend có thể không làm thay đổi schema phản hồi, ví dụ sai tổng tiền đơn hàng, sai trạng thái thanh toán, sai tồn kho, sai phân quyền hoặc side effect ghi xuống cơ sở dữ liệu không đúng. Do đó, hệ thống cần thêm các business assertion, kiểm tra invariant nghiệp vụ và kiểm tra trạng thái sau khi luồng kết thúc.",
        "Phạm vi kiểm thử cũng còn hẹp so với nhu cầu triển khai thực tế. Hệ thống hiện tập trung vào API/backend và chưa bao phủ đầy đủ kiểm thử hiệu năng, bảo mật, phân quyền nâng cao, contract giữa microservice, message queue/event-driven flow hoặc tương tác giao diện. Ngoài ra, phương pháp mới được kiểm chứng trên một hệ thống thương mại điện tử; với các miền nghiệp vụ như tài chính, viễn thông hoặc logistics, việc suy luận intent người dùng và thiết kế oracle có thể phức tạp hơn.",
        "Về vai trò AI, cách tiếp cận hiện tại vẫn thận trọng: AI hỗ trợ làm giàu dữ liệu nhưng không quyết định kết quả đạt/không đạt. Đây là lựa chọn phù hợp cho CI/CD vì bảo đảm tính tái lập. Tuy nhiên, hướng phát triển dài hạn có thể tiến tới một pipeline AI-driven hơn, trong đó AI chủ động phân cụm hành vi, suy luận intent, đề xuất assertion nghiệp vụ, sinh dữ liệu biên, tạo negative test, phân tích nguyên nhân fail và đề xuất cập nhật test khi API thay đổi. Điểm cần giữ lại là nguyên tắc AI-driven generation, deterministic validation: AI có thể điều phối và mở rộng kiểm thử, nhưng kết luận cuối cùng vẫn phải dựa trên oracle, rule hoặc bằng chứng có thể kiểm toán.",
    ]:
        add_paragraph(doc, text)
    doc.add_page_break()
    add_table(
        doc,
        ["Nhóm hạn chế", "Trạng thái hiện tại", "Hướng cải thiện"],
        [
            ["Nguồn dữ liệu", "Lưu lượng mô phỏng, log thử nghiệm; chưa phản ánh đầy đủ retry, timeout, request bất đồng bộ và hành vi đa dịch vụ.", "Kết nối log staging/production đã ẩn danh hóa; bổ sung quản trị dữ liệu, retention và kiểm toán truy cập."],
            ["Chất lượng log", "Giả định log có đủ session id, endpoint, status, ngữ cảnh xác thực và payload đã che dữ liệu.", "Chuẩn hóa schema log liên service, bổ sung correlation id, trace id và metadata nghiệp vụ cần thiết."],
            ["Golden oracle", "Dựa trên đặc tả API, lược đồ quan sát và quy tắc bỏ qua trường động.", "Mở rộng sang business assertions, invariant nghiệp vụ, kiểm tra side effect và trạng thái database/event sau luồng."],
            ["Độ phủ kiểm thử", "Tập trung vào API/backend; mutation testing còn 105 mutant inconclusive.", "Tăng baseline flow, sinh negative/boundary tests, mở rộng sang phân quyền, contract, event-driven flow và hiệu năng."],
            ["Khái quát hóa", "Mới kiểm chứng trên một hệ thống thương mại điện tử mã nguồn mở.", "Đánh giá thêm trên nhiều miền nghiệp vụ như tài chính, viễn thông, logistics hoặc hệ thống nội bộ doanh nghiệp."],
            ["AI-driven testing", "AI chủ yếu hỗ trợ tạo lưu lượng, đặt tên và gợi ý; chưa điều phối toàn bộ vòng đời kiểm thử.", "Tiến tới AI-driven generation: AI phân cụm hành vi, suy luận intent, đề xuất oracle, sinh test và phân tích fail; pass/fail vẫn được kiểm chứng tất định."],
            ["Vận hành CI/CD", "Pipeline có thể chạy lại nhưng chưa có đầy đủ cơ chế ưu tiên theo rủi ro và theo dõi xu hướng dài hạn.", "Tích hợp CI/CD, risk-based prioritization, cảnh báo hồi quy theo thời gian và human approval cho test/oracle mới."],
        ],
        widths=[1.35, 2.35, 2.7],
        caption="Bảng 6. Hạn chế hiện tại và hướng cải thiện",
    )


def add_conclusion_and_refs(doc):
    add_heading(doc, "IV. KẾT LUẬN", 1, page_break_before=True)
    for text in [
        "Mini-project đã xây dựng và chứng minh được một quy trình khép kín nhằm tự động hóa kiểm thử hồi quy backend/API dựa trên phân tích hành vi người dùng từ nhật ký hệ thống. Thay vì chỉ viết kiểm thử thủ công theo tài liệu đặc tả hoặc sinh kiểm thử trực tiếp từ mô tả bằng AI, hệ thống khai thác structured log như một nguồn dữ liệu thực nghiệm để tái dựng các chuỗi thao tác đã xuất hiện, chuẩn hóa chúng thành các luồng hành vi, sau đó sinh và chạy lại các bài kiểm thử API tương ứng.",
        "Về mặt kỹ thuật, đề tài đã hoàn thiện các thành phần chính của pipeline, bao gồm ghi nhận log có cấu trúc, che dữ liệu nhạy cảm, chuẩn hóa endpoint và phiên người dùng, khai phá luồng hành vi, sinh kiểm thử API, xây dựng golden oracle, chạy regression test, tổng hợp báo cáo và lưu trữ sản phẩm đầu ra. Các thành phần này được tổ chức tách bạch, có thể chạy lại theo từng bước và phù hợp với định hướng tích hợp vào quy trình CI/CD.",
        "Kết quả thử nghiệm cho thấy hướng tiếp cận là khả thi: hệ thống đã xử lý được dữ liệu mô phỏng, tái dựng các luồng hành vi ứng viên, tạo bộ kiểm thử ở mức API và phản ánh được trạng thái hồi quy qua các lượt chạy. Đánh giá đột biến cũng cung cấp tín hiệu rằng bộ kiểm thử có khả năng phát hiện một số nhóm lỗi liên quan đến định dạng và ràng buộc dữ liệu, dù vẫn cần mở rộng độ phủ để giảm số mutant inconclusive.",
        "Tuy vậy, phạm vi hiện tại cần được nhìn nhận thận trọng. Dữ liệu trong báo cáo là lưu lượng mô phỏng, hạ tầng triển khai ở quy mô cục bộ và đối tượng kiểm thử mới giới hạn trong một hệ thống thương mại điện tử mã nguồn mở. Khi áp dụng vào môi trường thật, hệ thống sẽ phải xử lý log nhiễu hơn, dữ liệu nhạy cảm phức tạp hơn, phiên người dùng đa dạng hơn, tương tác nhiều dịch vụ hơn và các ràng buộc nghiệp vụ thay đổi thường xuyên.",
        "Hướng phát triển tiếp theo không chỉ là tăng số lượng test được sinh ra, mà là nâng cấp pipeline thành một nền tảng kiểm thử hành vi có khả năng tự học, tự đề xuất và tự thích nghi. AI có thể đóng vai trò trung tâm trong quá trình khám phá luồng, suy luận intent, đề xuất assertion nghiệp vụ, sinh test biên và negative test, phân tích nguyên nhân thất bại, cũng như đề xuất cập nhật test khi API thay đổi. Đây là hướng tiếp cận AI-driven phù hợp để giảm công sức duy trì kiểm thử trong các hệ thống backend phát triển liên tục.",
        "Tuy nhiên, fully AI-driven không nên được hiểu là để AI tự quyết định hoàn toàn kết quả đúng sai mà không có cơ chế kiểm chứng. Trong kiểm thử hồi quy, đặc biệt khi tích hợp vào CI/CD, tính tất định, khả năng tái lập và khả năng giải thích vẫn là yêu cầu quan trọng. Vì vậy, hướng hợp lý là AI-driven generation, deterministic validation: AI chủ động hơn trong khám phá, sinh và phân tích kiểm thử, còn kết quả pass/fail cuối cùng vẫn được neo vào đặc tả API, business rules, invariant nghiệp vụ, oracle có thể kiểm toán và cơ chế phê duyệt của con người khi cần.",
        "Tổng kết lại, đóng góp chính của đề tài là hiện thực hóa một hướng tiếp cận kiểm thử hồi quy API dựa trên hành vi quan sát được từ log, kết hợp khai phá dữ liệu, sinh kiểm thử tự động và cơ chế oracle tất định. Hệ thống không thay thế hoàn toàn kiểm thử thủ công, nhưng có thể trở thành một lớp kiểm thử bổ sung quan trọng, giúp phát hiện hồi quy trên các luồng nghiệp vụ có nguồn gốc rõ ràng, giảm công sức bảo trì test case và nâng cao khả năng truy vết trong quá trình phát triển backend.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, page_break_before=True)
    refs = [
        "J. Pei, J. Han, B. Mortazavi-Asl, H. Pinto, Q. Chen, U. Dayal và M.-C. Hsu, \"PrefixSpan: Mining Sequential Patterns Efficiently by Prefix-Projected Pattern Growth,\" Proc. 17th IEEE Int. Conf. on Data Engineering (ICDE), 2001.",
        "R. Agrawal và R. Srikant, \"Mining Sequential Patterns,\" Proc. 11th IEEE Int. Conf. on Data Engineering (ICDE), 1995.",
        "R. Srikant và R. Agrawal, \"Mining Sequential Patterns: Generalizations and Performance Improvements,\" Proc. 5th Int. Conf. on Extending Database Technology (EDBT), 1996.",
        "Anthropic, \"Building Effective Agents,\" 2024, https://www.anthropic.com/engineering/building-effective-agents.",
        "Anthropic, \"Claude Agent SDK và Developer Documentation,\" https://docs.claude.com/.",
        "Medusa Documentation, https://docs.medusajs.com/.",
        "Playwright API Testing Documentation, https://playwright.dev/docs/api-testing.",
        "Elasticsearch, Logstash, Filebeat, Kibana Documentation, https://www.elastic.co/guide/.",
        "OpenAPI Specification, https://spec.openapis.org/oas/latest.html.",
        "PostgreSQL Documentation, https://www.postgresql.org/docs/.",
        "MinIO Documentation, https://min.io/docs/minio/.",
    ]
    for index, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"[{index}] {ref}", align="left")


def main():
    generate_report_figures()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_intro_chapter(doc)
    add_method_chapter(doc)
    add_results_chapter(doc)
    add_conclusion_and_refs(doc)
    doc.save(OUTPUT_FILE)
    print(f"Wrote {Path(OUTPUT_FILE).resolve()}")


if __name__ == "__main__":
    main()
